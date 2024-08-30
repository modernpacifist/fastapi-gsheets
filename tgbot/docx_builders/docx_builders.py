import os
import docx

from abc import ABC, abstractmethod
from datetime import datetime


class DocxBuilder(ABC):
    def __init__(self):
        self.path = './docx_files_storage'

    def save_doc(self, doc, name):
        try:
            if not os.path.exists(self.path):
                os.makedirs(self.path)

            path = f'{self.path}/{name}'
            doc.save(path)
            return path

        except Exception as e:
            print(f'Warning: could not create directory for docx files {e}')

    @abstractmethod
    def create(self):
        raise NotImplementedError('Method not implemented')


class ConferenceReport(DocxBuilder):
    def create(self, data, name):
        doc = docx.Document()

        doc.add_paragraph(f'Conference {data.get("id")} info:')

        doc.add_paragraph(
            f'Conference id: {data.get("id")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Google drive directory id: {data.get("google_drive_directory_id")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Full name: {data.get("name_rus")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Short name: {data.get("name_rus_short")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'English full name: {data.get("name_eng")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'English short name: {data.get("name_eng_short")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Organizator: {data.get("organized_by")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Registration start date: {data.get("registration_start_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Registration end date: {data.get("registration_end_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Submissions start date: {data.get("submission_start_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Submissions end date: {data.get("submission_end_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Conference starts at: {data.get("conf_start_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Conference ends at: {data.get("conf_end_date")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Organizer url: {data.get("url")}', style='List Bullet'
        )

        doc.add_paragraph(
            f'Organizer email: {data.get("email")}', style='List Bullet'
        )

        doc.add_page_break()

        return self.save_doc(doc, name)


class SingularReport(DocxBuilder):
    def create(self, data, name):
        doc = docx.Document()

        doc.add_paragraph('List:')
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'name'
        hdr_cells[1].text = 'submission time'
        hdr_cells[2].text = 'link'

        for d in data:
            row_cells = table.add_row().cells
            row_cells[0].text = d.get('name')
            dt = datetime.strptime(d.get('createdTime'), '%Y-%m-%dT%H:%M:%S.%fZ')
            row_cells[1].text = f'{dt.hour}:{dt.minute} {dt.day}.{dt.month}.{dt.year}'
            row_cells[2].text = d.get('webViewLink')

        doc.add_page_break()

        return self.save_doc(doc, name)
